"""
generate_word_doc.py
Generates the full API Automation Learning Guide as a formatted .docx file.
Run: py generate_word_doc.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── Colour Palette ────────────────────────────────────────────────────────────
C_DARK_BLUE   = RGBColor(0x1F, 0x49, 0x7D)
C_MED_BLUE    = RGBColor(0x2E, 0x74, 0xB5)
C_LIGHT_BLUE  = RGBColor(0xBD, 0xD7, 0xEE)
C_GREEN       = RGBColor(0x37, 0x86, 0x37)
C_RED         = RGBColor(0xA0, 0x00, 0x00)
C_ORANGE      = RGBColor(0xBF, 0x60, 0x00)
C_WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
C_GRAY        = RGBColor(0x59, 0x59, 0x59)
C_LIGHT_GRAY  = RGBColor(0xF2, 0xF2, 0xF2)


# ── XML helpers ───────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """Set the background fill colour of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_para_bg(para, hex_color):
    """Set the background shading of a paragraph."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)


def add_toc(doc):
    """Insert a Word auto-update Table of Contents field."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)


# ── Style helpers ─────────────────────────────────────────────────────────────

def h1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = C_DARK_BLUE
        run.font.size = Pt(18)
    return p


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = C_MED_BLUE
        run.font.size = Pt(14)
    return p


def h3(doc, text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = C_DARK_BLUE
        run.font.size = Pt(12)
    return p


def body(doc, text):
    p = doc.add_paragraph(text)
    p.style = doc.styles['Normal']
    for run in p.runs:
        run.font.size = Pt(11)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def code(doc, text):
    """Monospace code block with light-gray background."""
    for line in text.strip('\n').split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.right_indent = Inches(0.3)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        set_para_bg(p, 'F2F2F2')
        run = p.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = C_GRAY
    doc.add_paragraph()   # breathing room after code block


def note_box(doc, text, label="NOTE", color_hex='EAF3FB'):
    """A highlighted note/tip/info paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    set_para_bg(p, color_hex)
    label_run = p.add_run(f"{label}: ")
    label_run.bold = True
    label_run.font.color.rgb = C_DARK_BLUE
    label_run.font.size = Pt(11)
    body_run = p.add_run(text)
    body_run.font.size = Pt(11)
    doc.add_paragraph()


def table(doc, headers, rows, col_widths=None):
    """Add a formatted table with a dark-blue header row."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'

    # Header row
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_bg(hdr_cells[i], '1F497D')
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.font.color.rgb = C_WHITE
                run.font.bold = True
                run.font.size = Pt(10)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row_cells = t.rows[r_idx + 1].cells
        bg = 'FFFFFF' if r_idx % 2 == 0 else 'EBF3FB'
        for c_idx, cell_text in enumerate(row_data):
            row_cells[c_idx].text = cell_text
            set_cell_bg(row_cells[c_idx], bg)
            for para in row_cells[c_idx].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    # Column widths
    if col_widths:
        for row in t.rows:
            for i, cell in enumerate(row.cells):
                cell.width = Inches(col_widths[i])

    doc.add_paragraph()
    return t


def page_break(doc):
    doc.add_page_break()


def divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('─' * 80)
    run.font.color.rgb = C_LIGHT_BLUE
    run.font.size = Pt(8)


# ── COVER PAGE ────────────────────────────────────────────────────────────────

def cover_page(doc):
    for _ in range(6):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run('API TEST AUTOMATION')
    title_run.font.size = Pt(32)
    title_run.font.bold = True
    title_run.font.color.rgb = C_DARK_BLUE

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run('Complete Learning Guide')
    sub_run.font.size = Pt(20)
    sub_run.font.color.rgb = C_MED_BLUE

    doc.add_paragraph()
    line_p = doc.add_paragraph()
    line_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line_run = line_p.add_run('─' * 50)
    line_run.font.color.rgb = C_MED_BLUE

    doc.add_paragraph()
    api_p = doc.add_paragraph()
    api_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    api_run = api_p.add_run('Simple Books API  ·  Java REST Assured  ·  Python pytest  ·  TypeScript Playwright')
    api_run.font.size = Pt(13)
    api_run.font.color.rgb = C_GRAY

    for _ in range(8):
        doc.add_paragraph()

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_p.add_run(f'Generated: {datetime.date.today().strftime("%B %d, %Y")}')
    date_run.font.size = Pt(11)
    date_run.font.color.rgb = C_GRAY

    page_break(doc)


# ── CHAPTER 1 ─────────────────────────────────────────────────────────────────

def chapter_1(doc):
    h1(doc, 'Chapter 1 — What is an API?')

    h2(doc, '1.1 The Simple Explanation')
    body(doc, 'API stands for Application Programming Interface.')
    body(doc, 'Think of it like a waiter in a restaurant:')
    bullet(doc, 'You (the client) sit at the table and want food.')
    bullet(doc, 'The kitchen (the server) has the food.')
    bullet(doc, 'You don\'t go into the kitchen yourself — the waiter (API) takes your order and brings back the food.')
    body(doc, 'In software:')
    bullet(doc, 'Your app (client) needs data or wants to do something.')
    bullet(doc, 'A database or service (server) has that data.')
    bullet(doc, 'The API is the agreed way for your app to ask for it and get a response.')

    h2(doc, '1.2 Real-World Examples')
    table(doc,
        ['Service', 'What the API Does'],
        [
            ['Google Maps', 'Your app asks for directions → API returns a route'],
            ['PayPal', 'Your shop sends payment request → API confirms or rejects'],
            ['Weather App', 'Your phone asks current temp → API returns weather data'],
            ['Simple Books API', 'Your test asks for books → API returns a JSON list of books'],
        ],
        col_widths=[2.0, 4.2]
    )

    h2(doc, '1.3 How an API Call Works (Step by Step)')
    code(doc, """\
YOUR PROGRAM                        SERVER (API)
     |                                   |
     |  1. Send HTTP Request  -------->  |
     |     Method : GET                  |
     |     URL    : /books               |
     |     Headers: Content-Type: json   |
     |                                   |
     |  <--------  2. Receive Response   |
     |     Status : 200 OK               |
     |     Body   : [{"id":1, ...}]      |
     |                                   |""")
    note_box(doc,
        'The client NEVER touches the server directly — it only communicates through the API '
        'using well-defined HTTP requests and responses.',
        label='KEY CONCEPT')

    page_break(doc)


# ── CHAPTER 2 ─────────────────────────────────────────────────────────────────

def chapter_2(doc):
    h1(doc, 'Chapter 2 — What is REST?')
    body(doc, 'REST = REpresentational State Transfer. It is a set of design rules for APIs. '
              'An API that follows REST rules is called a RESTful API.')

    h2(doc, '2.1 The 4 Core REST Rules')

    h3(doc, 'Rule 1 — Resources are identified by URLs')
    body(doc, 'Every piece of data has its own unique address (URL):')
    code(doc, """\
https://simple-books-api.click/books        ← all books
https://simple-books-api.click/books/1      ← one specific book (ID = 1)
https://simple-books-api.click/orders       ← all orders
https://simple-books-api.click/orders/abc   ← one specific order""")

    h3(doc, 'Rule 2 — Use HTTP Methods to say WHAT you want to do')
    table(doc,
        ['HTTP Method', 'Meaning', 'Example'],
        [
            ['GET',    'Read / Fetch existing data',   'Get list of books'],
            ['POST',   'Create a new resource',        'Place a new order'],
            ['PUT',    'Replace a resource completely', 'Replace entire order'],
            ['PATCH',  'Partially update a resource',  'Change just the customer name'],
            ['DELETE', 'Remove a resource',            'Cancel an order'],
        ],
        col_widths=[1.4, 2.2, 2.8]
    )

    h3(doc, 'Rule 3 — Stateless')
    body(doc, 'The server does NOT remember your previous request. Every request must carry '
              'all the information it needs, including the authentication token.')

    h3(doc, 'Rule 4 — Use standard HTTP status codes')
    body(doc, 'The server communicates the result using a 3-digit number (status code). '
              'See Chapter 3 for the full reference.')

    page_break(doc)


# ── CHAPTER 3 ─────────────────────────────────────────────────────────────────

def chapter_3(doc):
    h1(doc, 'Chapter 3 — HTTP Fundamentals')

    h2(doc, '3.1 The Request — What You Send')
    body(doc, 'Every HTTP request has these parts:')
    code(doc, """\
┌─────────────────────────────────────────────────────────┐
│  REQUEST LINE                                           │
│  POST https://simple-books-api.click/orders HTTP/1.1    │
│                                                         │
│  HEADERS  (metadata about the request)                  │
│  Content-Type: application/json                         │
│  Authorization: Bearer eyJhbGciOiJIUzI1NiIsInR...      │
│                                                         │
│  BODY  (the data you are sending — only for POST/PATCH) │
│  {                                                      │
│    "bookId": 1,                                         │
│    "customerName": "John Doe"                           │
│  }                                                      │
└─────────────────────────────────────────────────────────┘""")

    table(doc,
        ['Part', 'Purpose', 'Example'],
        [
            ['Method',  'What action to take',                  'GET, POST, PATCH, DELETE'],
            ['URL',     'Which resource to act on',             'https://...click/orders'],
            ['Headers', 'Extra info (auth, content type)',      'Authorization: Bearer abc'],
            ['Body',    'Data sent to the server (POST/PATCH)', '{"bookId": 1, "customerName": "John"}'],
        ],
        col_widths=[1.2, 2.2, 3.0]
    )

    note_box(doc, 'GET and DELETE requests usually have NO body. POST and PATCH usually have a body.')

    h2(doc, '3.2 The Response — What You Receive')
    code(doc, """\
┌─────────────────────────────────────────────────────────┐
│  STATUS LINE                                            │
│  HTTP/1.1 201 Created                                   │
│                                                         │
│  HEADERS                                                │
│  Content-Type: application/json                         │
│                                                         │
│  BODY  (the data returned by the server)                │
│  {                                                      │
│    "created": true,                                     │
│    "orderId": "PF6MflPDcuhWobZcgmJy5"                   │
│  }                                                      │
└─────────────────────────────────────────────────────────┘""")

    h2(doc, '3.3 HTTP Status Codes')
    body(doc, 'The first digit tells you the category:')
    code(doc, """\
1xx → Informational  (rarely seen in testing)
2xx → SUCCESS        ← you want these
3xx → Redirect       (resource moved somewhere else)
4xx → CLIENT ERROR   ← your request was wrong
5xx → SERVER ERROR   ← the server crashed""")

    h3(doc, 'Status Codes You Will See in These Tests')
    table(doc,
        ['Code', 'Name', 'Meaning', 'When You See It'],
        [
            ['200', 'OK',                   'Request succeeded, body has data',        'Successful GET requests'],
            ['201', 'Created',              'New resource was created',                'POST /orders success'],
            ['204', 'No Content',           'Success but no response body',            'PATCH and DELETE success'],
            ['400', 'Bad Request',          'You sent invalid or missing data',        'Missing customerName'],
            ['401', 'Unauthorized',         'No token or invalid token',               'Calling /orders without auth'],
            ['403', 'Forbidden',            'Token valid but no permission',           'Accessing another client\'s order'],
            ['404', 'Not Found',            'Resource doesn\'t exist',                'GET /books/99999'],
            ['409', 'Conflict',             'Duplicate resource',                      'Same email registered twice'],
            ['500', 'Internal Server Error','Server crashed or bug on server side',    'Not your fault'],
        ],
        col_widths=[0.6, 1.6, 2.4, 1.8]
    )

    h2(doc, '3.4 Common Request Headers')
    table(doc,
        ['Header', 'Value', 'Meaning'],
        [
            ['Content-Type',   'application/json',    'The body I\'m sending is JSON format'],
            ['Authorization',  'Bearer <token>',      'Proves you are authenticated'],
            ['Accept',         'application/json',    'Please respond with JSON format'],
        ],
        col_widths=[1.8, 2.0, 2.6]
    )

    h2(doc, '3.5 JSON — The Data Format')
    body(doc, 'Almost all REST APIs use JSON (JavaScript Object Notation). It has two structures:')

    h3(doc, 'Object — key/value pairs, surrounded by { }')
    code(doc, """\
{
  "id": 1,
  "name": "The Russian",
  "type": "fiction",
  "available": true
}""")

    h3(doc, 'Array — an ordered list, surrounded by [ ]')
    code(doc, """\
[
  { "id": 1, "name": "The Russian" },
  { "id": 2, "name": "Just as I Am" },
  { "id": 3, "name": "The Vanishing Half" }
]""")

    table(doc,
        ['JSON Type', 'Example', 'Notes'],
        [
            ['String',  '"John Doe"',         'Always in double quotes'],
            ['Number',  '1  or  3.14',        'No quotes'],
            ['Boolean', 'true  or  false',    'Lowercase, no quotes'],
            ['Null',    'null',               'Represents "no value"'],
            ['Object',  '{ "key": "value" }', 'Key-value pairs in curly braces'],
            ['Array',   '[1, 2, 3]',          'Ordered list in square brackets'],
        ],
        col_widths=[1.2, 1.8, 3.4]
    )

    page_break(doc)


# ── CHAPTER 4 ─────────────────────────────────────────────────────────────────

def chapter_4(doc):
    h1(doc, 'Chapter 4 — Authentication — Bearer Tokens')

    h2(doc, '4.1 Why Authentication is Needed')
    body(doc, 'Without authentication, anyone could:')
    bullet(doc, 'See your private orders')
    bullet(doc, 'Delete your orders')
    bullet(doc, 'Create fake orders in your name')
    body(doc, 'Authentication proves WHO you are so the API only gives you YOUR data.')

    h2(doc, '4.2 How Bearer Token Authentication Works')
    code(doc, """\
STEP 1: Register as an API client
──────────────────────────────────────────────────────────
You  →  POST /api-clients/
        Body: { "clientName": "MyApp", "clientEmail": "me@example.com" }

API  →  201 Created
        Body: { "accessToken": "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9..." }

        ↑ This long string is your TOKEN. Copy it and store it safely.

STEP 2: Use the token in every protected request
──────────────────────────────────────────────────────────
You  →  POST /orders
        Header: Authorization: Bearer eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9...
        Body  : { "bookId": 1, "customerName": "John" }

API  →  201 Created  ✓  (token is valid)

Without the token:
You  →  POST /orders  (no Authorization header)
API  →  401 Unauthorized  ✗""")

    h2(doc, '4.3 What Does "Bearer" Mean?')
    body(doc, 'The word Bearer means "the person holding this token". The format is:')
    code(doc, 'Authorization: Bearer <your-token-here>')
    note_box(doc, 'The word "Bearer" is REQUIRED. If you just send the token without it, '
                  'the API will return 401 Unauthorized.', label='IMPORTANT')

    h2(doc, '4.4 Why We Use a Unique Email Every Run')
    body(doc, 'If you try to register the same email twice, the API returns 409 Conflict. '
              'To avoid this in automated tests, we generate a random suffix every run:')
    code(doc, """\
// Java
String uniqueId = UUID.randomUUID().toString().substring(0, 8);
// e.g. "a3f7c2d1" — different every run

# Python
unique_id = str(uuid.uuid4())[:8]

// TypeScript
const uniqueId = Date.now().toString(36);  // e.g. "lxyz123" """)

    page_break(doc)


# ── CHAPTER 5 ─────────────────────────────────────────────────────────────────

def chapter_5(doc):
    h1(doc, 'Chapter 5 — What is API Testing?')

    h2(doc, '5.1 Definition')
    body(doc, 'API testing means sending HTTP requests to an API and verifying the '
              'responses are correct — the right status code, the right data, '
              'the right error messages.')
    body(doc, 'Unlike UI testing (which clicks buttons), API testing:')
    bullet(doc, 'Talks directly to the server — no browser needed')
    bullet(doc, 'Is faster — no rendering or DOM loading')
    bullet(doc, 'Is more stable — UI changes don\'t break it')
    bullet(doc, 'Tests the actual business logic at the core')

    h2(doc, '5.2 What We Check in Every Test')
    table(doc,
        ['What We Check', 'Example Assertion'],
        [
            ['Status code',         'response.statusCode == 201'],
            ['Response body field', 'body.created == true'],
            ['Field value',         'body.customerName == "John Doe"'],
            ['Field presence',      'body has an orderId key'],
            ['Data type',           'body.id is a number'],
            ['Error message',       'body.error exists when 400'],
        ],
        col_widths=[2.0, 4.4]
    )

    h2(doc, '5.3 Types of Tests')

    h3(doc, 'Happy Path Tests')
    body(doc, 'The normal case — valid data, correct auth, expected to succeed.')
    code(doc, 'POST /orders with valid bookId + customerName  →  201 Created  ✓')

    h3(doc, 'Negative Tests')
    body(doc, 'Deliberately wrong input — expected to fail gracefully.')
    code(doc, """\
POST /orders  with no auth token           →  401 Unauthorized
POST /orders  missing customerName         →  400 Bad Request
GET  /books/99999  (doesn't exist)         →  404 Not Found
GET  /books?type=magazine  (invalid type)  →  400 Bad Request""")

    h3(doc, 'Boundary Tests')
    body(doc, 'Testing the edges of allowed input ranges.')
    code(doc, """\
GET /books?limit=1   →  minimum allowed  ✓
GET /books?limit=20  →  maximum allowed  ✓
GET /books?limit=0   →  below minimum   →  400  ✗
GET /books?limit=21  →  above maximum   →  400  ✗""")

    h3(doc, 'CRUD Lifecycle Test')
    body(doc, 'Testing the full Create → Read → Update → Delete flow in one sequence:')
    code(doc, """\
Step 1:  POST   /orders              →  201  (save the orderId)
Step 2:  GET    /orders/{orderId}    →  200  (verify data matches what we sent)
Step 3:  PATCH  /orders/{orderId}    →  204  (update the name)
Step 4:  GET    /orders/{orderId}    →  200  (verify the update was saved)
Step 5:  DELETE /orders/{orderId}    →  204  (delete it)
Step 6:  GET    /orders/{orderId}    →  404  (confirm it's gone)""")

    h2(doc, '5.4 Test Coverage Checklist')
    body(doc, 'For every API endpoint you test, ask these questions:')
    bullet(doc, 'Does the happy path return the correct status and body?')
    bullet(doc, 'What happens with missing required fields? (expect 400)')
    bullet(doc, 'What happens with an invalid resource ID? (expect 404)')
    bullet(doc, 'What happens without authentication? (expect 401)')
    bullet(doc, 'What happens with invalid parameter values? (expect 400)')
    bullet(doc, 'Does CRUD correctly change and persist data?')

    page_break(doc)


# ── CHAPTER 6 ─────────────────────────────────────────────────────────────────

def chapter_6(doc):
    h1(doc, 'Chapter 6 — The Simple Books API — All Endpoints')

    body(doc, 'Base URL:  https://simple-books-api.click')
    doc.add_paragraph()

    table(doc,
        ['Method', 'Endpoint', 'Auth?', 'Purpose'],
        [
            ['GET',    '/status',            'No',  'API health check'],
            ['GET',    '/books',             'No',  'List books (filter by type/limit)'],
            ['GET',    '/books/{bookId}',    'No',  'Get full details of one book'],
            ['POST',   '/api-clients/',      'No',  'Register and receive access token'],
            ['POST',   '/orders',            'YES', 'Place a new order'],
            ['GET',    '/orders',            'YES', 'List all your orders'],
            ['GET',    '/orders/{orderId}',  'YES', 'Get one specific order'],
            ['PATCH',  '/orders/{orderId}',  'YES', 'Update order customer name'],
            ['DELETE', '/orders/{orderId}',  'YES', 'Delete an order'],
        ],
        col_widths=[0.8, 1.9, 0.7, 3.0]
    )

    h2(doc, '6.1 GET /status')
    code(doc, """\
Request:
  GET https://simple-books-api.click/status

Response (200 OK):
  { "status": "OK" }""")

    h2(doc, '6.2 GET /books')
    body(doc, 'Query parameters (optional, added after ? in the URL):')
    table(doc,
        ['Parameter', 'Valid Values', 'Example URL'],
        [
            ['type',  'fiction  or  non-fiction', '/books?type=fiction'],
            ['limit', '1 to 20 (integer)',         '/books?limit=5'],
        ],
        col_widths=[1.2, 2.2, 2.8]
    )
    code(doc, """\
Request:
  GET https://simple-books-api.click/books?type=fiction&limit=3

Response (200 OK):
  [
    { "id": 1, "name": "The Russian",       "type": "fiction", "available": true  },
    { "id": 3, "name": "The Vanishing Half","type": "fiction", "available": true  }
  ]""")

    h2(doc, '6.3 GET /books/{bookId}')
    code(doc, """\
Request:
  GET https://simple-books-api.click/books/1

Response (200 OK):
  {
    "id": 1,
    "name": "The Russian",
    "author": "James Patterson and James O. Born",
    "isbn": "1780899475",
    "type": "fiction",
    "price": 12.98,
    "current-stock": 12,
    "available": true
  }

Error response:
  GET /books/99999  →  404 Not Found
  { "error": "No book with id 99999" }""")

    h2(doc, '6.4 POST /api-clients/ — Get Your Token')
    code(doc, """\
Request:
  POST https://simple-books-api.click/api-clients/
  Content-Type: application/json

  {
    "clientName": "MyTestApp",
    "clientEmail": "myapp@example.com"
  }

Response (201 Created):
  { "accessToken": "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9..." }

If same email used again:
  409 Conflict
  { "error": "API client already registered. Try a different email." }""")

    h2(doc, '6.5 POST /orders — Create an Order')
    code(doc, """\
Request:
  POST https://simple-books-api.click/orders
  Content-Type: application/json
  Authorization: Bearer <your-token>

  {
    "bookId": 1,
    "customerName": "John Doe"
  }

Response (201 Created):
  {
    "created": true,
    "orderId": "PF6MflPDcuhWobZcgmJy5"
  }

Without token:   401 Unauthorized  →  { "error": "Missing Authorization header." }
Missing field:   400 Bad Request   →  { "error": "customerName" }""")

    h2(doc, '6.6 PATCH /orders/{orderId} — Update Order')
    code(doc, """\
Request:
  PATCH https://simple-books-api.click/orders/PF6MflPDcuhWobZcgmJy5
  Content-Type: application/json
  Authorization: Bearer <your-token>

  { "customerName": "Jane Smith" }

Response:  204 No Content  (success, but NO body is returned)""")
    note_box(doc, '204 No Content is a SUCCESS response. Do NOT call response.json() '
                  'on a 204 — there is no body and it will throw an error.', label='IMPORTANT')

    h2(doc, '6.7 DELETE /orders/{orderId}')
    code(doc, """\
Request:
  DELETE https://simple-books-api.click/orders/PF6MflPDcuhWobZcgmJy5
  Authorization: Bearer <your-token>

Response:  204 No Content  (deleted, no body)""")

    page_break(doc)


# ── CHAPTER 7 ─────────────────────────────────────────────────────────────────

def chapter_7(doc):
    h1(doc, 'Chapter 7 — Java — REST Assured + TestNG')

    h2(doc, '7.1 What is REST Assured?')
    body(doc, 'REST Assured is a Java library that makes it very easy to test REST APIs. '
              'It uses a fluent "given / when / then" style — the code reads like plain English:')
    code(doc, """\
given()
    .header("Authorization", "Bearer " + token)
    .body(orderData)
.when()
    .post("/orders")
.then()
    .statusCode(201)
    .body("created", equalTo(true));

// Reading this aloud:
// "Given a header and body, when I POST to /orders,
//  then the status code should be 201 and 'created' should be true." """)

    h2(doc, '7.2 What is TestNG?')
    body(doc, 'TestNG is a test framework for Java. It organises and runs tests, '
              'controls execution order, and supports parallel execution.')

    h3(doc, 'Key TestNG Annotations')
    table(doc,
        ['Annotation', 'When It Runs', 'Used For'],
        [
            ['@BeforeSuite',  'Once, before ALL tests',             'One-time global setup'],
            ['@BeforeClass',  'Once, before tests in this CLASS',   'Get auth token for the class'],
            ['@BeforeMethod', 'Before EACH test method',            'Per-test setup (e.g., clear data)'],
            ['@Test',         'Marks a method as a test case',      'Every test you write'],
            ['@AfterMethod',  'After each test method',             'Per-test cleanup'],
            ['@AfterClass',   'After all tests in the class',       'Class-level cleanup'],
            ['@AfterSuite',   'Once, after ALL tests end',          'Final cleanup (close connections)'],
        ],
        col_widths=[1.5, 2.2, 2.7]
    )

    h2(doc, '7.3 Project Structure')
    code(doc, """\
java-rest-assured/
├── pom.xml                           ← Maven (library manager)
├── testng.xml                        ← Test suite definition (parallel settings)
└── src/test/java/com/simplebooksapi/
    ├── base/
    │   └── BaseTest.java             ← Shared base URL + request config
    ├── utils/
    │   └── AuthHelper.java           ← Token registration and caching
    └── tests/
        ├── StatusTest.java           ← Tests for GET /status
        ├── BooksTest.java            ← Tests for GET /books
        └── OrdersTest.java           ← Full CRUD order tests""")

    h2(doc, '7.4 pom.xml — Dependency Manager')
    body(doc, 'Maven\'s pom.xml is your project\'s "shopping list". Instead of manually '
              'downloading JAR files, you declare what you need and Maven downloads it automatically.')
    code(doc, """\
<!-- REST Assured: sends HTTP requests and validates responses -->
<dependency>
    <groupId>io.rest-assured</groupId>
    <artifactId>rest-assured</artifactId>
    <version>5.4.0</version>
    <scope>test</scope>
</dependency>

<!-- TestNG: the test runner and report generator -->
<dependency>
    <groupId>org.testng</groupId>
    <artifactId>testng</artifactId>
    <version>7.9.0</version>
    <scope>test</scope>
</dependency>

<!-- Jackson: converts Java Maps into JSON strings automatically -->
<dependency>
    <groupId>com.fasterxml.jackson.core</groupId>
    <artifactId>jackson-databind</artifactId>
    <version>2.16.1</version>
</dependency>""")

    h2(doc, '7.5 BaseTest.java — The Foundation')
    body(doc, 'Every test class extends BaseTest. It configures the base URL, '
              'content type, and console logging ONCE so all tests inherit it.')
    code(doc, """\
public class BaseTest {
    public static final String BASE_URL = "https://simple-books-api.click";
    protected static RequestSpecification requestSpec;

    @BeforeSuite
    public void setupSuite() {
        requestSpec = new RequestSpecBuilder()
            .setBaseUri(BASE_URL)                    // All calls go to this root URL
            .setContentType("application/json")      // Content-Type: application/json
            .addFilter(new RequestLoggingFilter())   // Print outgoing requests
            .addFilter(new ResponseLoggingFilter())  // Print incoming responses
            .build();
    }
}""")

    h2(doc, '7.6 AuthHelper.java — Token Management')
    body(doc, 'Handles registering an API client and caching the token. '
              'Uses a static field so registration happens only ONCE per test run.')
    code(doc, """\
public class AuthHelper {
    private static String accessToken;  // static = shared across all tests

    public static String getAccessToken(RequestSpecification spec) {
        if (accessToken != null) return accessToken;  // return cached token

        String uniqueId = UUID.randomUUID().toString().substring(0, 8);
        Map<String, String> body = new HashMap<>();
        body.put("clientName",  "TestClient_" + uniqueId);
        body.put("clientEmail", "testclient_" + uniqueId + "@example.com");

        Response response = RestAssured.given()
                .spec(spec).body(body).post("/api-clients/");

        response.then().statusCode(201);
        accessToken = response.jsonPath().getString("accessToken");
        return accessToken;
    }
}""")

    h2(doc, '7.7 Understanding JSON Path in REST Assured')
    body(doc, 'jsonPath() lets you extract values from the JSON response body:')
    code(doc, """\
// Response body:
// { "order": { "id": "abc", "book": { "name": "The Russian" } } }

response.jsonPath().getString("order.id")          // returns "abc"
response.jsonPath().getString("order.book.name")   // returns "The Russian"
response.jsonPath().getInt("order.bookId")         // returns integer
response.jsonPath().getBoolean("created")          // returns boolean""")

    h2(doc, '7.8 Query Parameters vs Path Parameters')
    code(doc, """\
// QUERY PARAMETER → added after ? in the URL
// Result URL: /books?type=fiction
.queryParam("type", "fiction")

// PATH PARAMETER → replaces a {placeholder} in the URL
// Result URL: /books/1
.pathParam("bookId", 1)
.get("/books/{bookId}")    ← {bookId} is replaced with 1""")

    h2(doc, '7.9 OrdersTest.java — Key Concepts')
    body(doc, 'The orders test chain uses dependsOnMethods to ensure proper ordering:')
    code(doc, """\
@Test(description = "Create order")
public void testCreateOrder() {
    // ... creates order
    orderId = response.jsonPath().getString("orderId");
}

@Test(dependsOnMethods = "testCreateOrder")  // only runs if Create passed
public void testGetOrder() {
    // orderId is now available because testCreateOrder ran first
}

@Test(dependsOnMethods = "testGetOrder")
public void testUpdateOrder() { ... }

@Test(dependsOnMethods = "testUpdateOrder")
public void testDeleteOrder() { ... }""")

    h2(doc, '7.10 How to Run — Java')
    table(doc,
        ['Command', 'What It Does'],
        [
            ['mvn test',                               'Run all tests in the suite'],
            ['mvn test -Dtest=BooksTest',              'Run only BooksTest class'],
            ['mvn test -Dtest=BooksTest#testGetAll',   'Run one specific test method'],
        ],
        col_widths=[3.2, 3.2]
    )

    page_break(doc)


# ── CHAPTER 8 ─────────────────────────────────────────────────────────────────

def chapter_8(doc):
    h1(doc, 'Chapter 8 — Python — pytest + requests')

    h2(doc, '8.1 What is requests?')
    body(doc, 'requests is a Python library for sending HTTP requests. '
              'It is one of the most downloaded Python packages ever.')
    code(doc, """\
import requests

# GET request
response = requests.get("https://simple-books-api.click/books")

# POST request with JSON body
response = requests.post(
    "https://simple-books-api.click/orders",
    json={"bookId": 1, "customerName": "John"},   # json= auto-sets Content-Type
    headers={"Authorization": "Bearer abc123"}
)

# Reading the response
print(response.status_code)    # 201
print(response.json())         # {'created': True, 'orderId': 'abc123'}
print(response.text)           # raw JSON string
print(response.ok)             # True  (status < 400)""")

    h2(doc, '8.2 What is pytest?')
    body(doc, 'pytest is Python\'s most popular test framework. It automatically finds '
              'and runs test files and provides a powerful setup/teardown system called fixtures.')

    h3(doc, 'pytest Discovery Rules')
    code(doc, """\
tests/
├── test_status.py    ← FOUND   (file starts with test_)
├── test_books.py     ← FOUND
└── helper.py         ← IGNORED (doesn't start with test_)

Inside a file:
  class TestBooks:                 ← FOUND   (class starts with Test)
      def test_get_books(self):    ← FOUND   (method starts with test_)
      def helper_method(self):     ← IGNORED""")

    h2(doc, '8.3 Fixtures — The Most Important pytest Concept')
    body(doc, 'A fixture is a reusable setup function. You write it once with @pytest.fixture '
              'and any test that needs it simply adds it as a parameter — pytest injects it automatically.')

    h3(doc, 'Without fixtures — repetitive and messy')
    code(doc, """\
def test_get_books():
    BASE_URL = "https://simple-books-api.click"  # repeated in every test
    response = requests.get(f"{BASE_URL}/books")
    assert response.status_code == 200

def test_get_status():
    BASE_URL = "https://simple-books-api.click"  # repeated again
    response = requests.get(f"{BASE_URL}/status")""")

    h3(doc, 'With fixtures — clean and reusable')
    code(doc, """\
# conftest.py
@pytest.fixture(scope="session")
def base_url():
    return "https://simple-books-api.click"

# test_books.py
def test_get_books(base_url):    # pytest sees 'base_url' and injects the value
    response = requests.get(f"{base_url}/books")
    assert response.status_code == 200

def test_get_status(base_url):   # same fixture, reused — no duplication
    response = requests.get(f"{base_url}/status")""")

    h2(doc, '8.4 Fixture Scope — The Lifecycle')
    table(doc,
        ['scope=', 'Fixture Runs...', 'Use For'],
        [
            ['"function"', 'Before every single test (default)', 'Setup that must be fresh each time'],
            ['"class"',    'Once per TestClass',                  'Shared state for one class'],
            ['"module"',   'Once per .py file',                   'Per-file shared resources'],
            ['"session"',  'Once for the entire test run',        'Token, database connection (expensive setup)'],
        ],
        col_widths=[1.2, 2.4, 2.8]
    )
    note_box(doc,
        'scope="session" for the auth token means we register the API client ONCE, '
        'not before every test. Without this, 9 tests would create 9 different API clients.',
        label='WHY IT MATTERS')

    h2(doc, '8.5 conftest.py — The Special File')
    body(doc, 'conftest.py is automatically loaded by pytest. Fixtures defined here are '
              'available to all test files in the same folder and subfolders without any imports.')
    code(doc, """\
# conftest.py
import pytest, requests, uuid

@pytest.fixture(scope="session")
def base_url():
    return "https://simple-books-api.click"

@pytest.fixture(scope="session")
def auth_token(base_url):              # this fixture uses the base_url fixture
    unique_id = str(uuid.uuid4())[:8]
    payload = {
        "clientName":  f"TestClient_{unique_id}",
        "clientEmail": f"testclient_{unique_id}@example.com"
    }
    response = requests.post(f"{base_url}/api-clients/", json=payload)
    assert response.status_code == 201, f"Registration failed: {response.text}"
    return response.json()["accessToken"]

@pytest.fixture(scope="session")
def auth_headers(auth_token):          # this fixture uses auth_token fixture
    return {"Authorization": f"Bearer {auth_token}"}""")

    h2(doc, '8.6 test_orders.py — Sharing State Between Test Methods')
    body(doc, 'Tests in a class run in definition order. We use a class variable '
              'to pass the orderId from the Create test to all later tests:')
    code(doc, """\
class TestOrders:
    order_id = None    # class variable: shared across all methods in this class

    def test_01_create_order(self, base_url, auth_headers):
        response = requests.post(f"{base_url}/orders",
                                 json={"bookId":1, "customerName":"John Doe"},
                                 headers=auth_headers)
        assert response.status_code == 201
        TestOrders.order_id = response.json()["orderId"]  # SAVE for later tests

    def test_03_get_order_by_id(self, base_url, auth_headers):
        # Uses the order_id saved by test_01
        response = requests.get(f"{base_url}/orders/{TestOrders.order_id}",
                                headers=auth_headers)
        assert response.status_code == 200
        assert response.json()["customerName"] == "John Doe"

    def test_04_update_order(self, base_url, auth_headers):
        response = requests.patch(f"{base_url}/orders/{TestOrders.order_id}",
                                  json={"customerName": "Jane Smith"},
                                  headers=auth_headers)
        assert response.status_code == 204
        # DO NOT call response.json() on a 204 — there is no body!""")

    h2(doc, '8.7 How to Run — Python')
    table(doc,
        ['Command', 'What It Does'],
        [
            ['pip install -r requirements.txt',            'Install all dependencies'],
            ['pytest',                                     'Run all tests'],
            ['pytest -v',                                  'Verbose: show each test name'],
            ['pytest tests/test_orders.py',               'Run one file'],
            ['pytest tests/test_orders.py::TestOrders::test_01_create_order', 'Run one method'],
            ['pytest --html=report.html',                 'Generate HTML report'],
            ['pytest -x',                                  'Stop on first failure'],
            ['pytest -k "create or delete"',              'Run tests matching keyword'],
        ],
        col_widths=[3.8, 2.6]
    )

    page_break(doc)


# ── CHAPTER 9 ─────────────────────────────────────────────────────────────────

def chapter_9(doc):
    h1(doc, 'Chapter 9 — TypeScript — Playwright API Tests')

    h2(doc, '9.1 What is Playwright?')
    body(doc, 'Playwright is a modern automation framework built by Microsoft. '
              'Most people use it for browser testing, but it also has a powerful '
              'built-in HTTP client for API testing — no extra library needed.')
    bullet(doc, 'No need to install axios or node-fetch')
    bullet(doc, 'Same tool covers both UI and API tests')
    bullet(doc, 'Built-in HTML reporter with visual results')
    bullet(doc, 'Native TypeScript support')
    bullet(doc, 'Parallel test execution out of the box')

    h2(doc, '9.2 TypeScript Basics for Beginners')
    body(doc, 'TypeScript is JavaScript with type annotations — you declare what type '
              'a variable holds, and mistakes are caught before the code even runs.')
    code(doc, """\
// JavaScript — no types, errors only appear at runtime:
let token = "abc123";
let orderId;           // could be anything — dangerous

// TypeScript — types declared, errors caught at compile time:
let token: string = "abc123";   // MUST always be a string
let orderId: string;             // will be assigned later, but must be a string

// TypeScript prevents this mistake:
token = 42;   // ERROR: Type 'number' is not assignable to type 'string'""")

    h2(doc, '9.3 Async/Await — Handling HTTP Requests')
    body(doc, 'All HTTP calls take time (the network is slow compared to code). '
              'We use async/await to tell JavaScript: "wait here until the response arrives".')
    code(doc, """\
// WITHOUT await — WRONG:
const response = request.get('/status');    // response is a Promise, not the real response
console.log(response.status());             // ERROR: Promise has no .status() method

// WITH await — CORRECT:
const response = await request.get('/status');  // waits for the HTTP call to finish
console.log(response.status());                  // 200  ✓

// Any function using 'await' MUST be declared as 'async':
test('my test', async ({ request }) => {   // 'async' keyword is required
    const response = await request.get('/status');
    expect(response.status()).toBe(200);
});""")

    h2(doc, '9.4 The request Fixture')
    body(doc, 'Playwright provides a built-in HTTP client called request. '
              'You just declare it as a parameter and Playwright injects it automatically — '
              'just like pytest fixtures:')
    code(doc, """\
test('my test', async ({ request }) => {
    //                    ↑
    // Playwright provides this — you never instantiate it yourself.
    // It already knows the baseURL from playwright.config.ts.

    const response = await request.get('/status');
    expect(response.status()).toBe(200);
});""")

    h2(doc, '9.5 Making Requests in Playwright')
    code(doc, """\
// GET (simple)
const r = await request.get('/books');

// GET with query params → /books?type=fiction&limit=3
const r = await request.get('/books', {
    params: { type: 'fiction', limit: 3 }
});

// POST with body and auth header
const r = await request.post('/orders', {
    headers: { Authorization: `Bearer ${token}` },
    data: { bookId: 1, customerName: 'John Doe' }  // auto-serialised to JSON
});

// PATCH
const r = await request.patch(`/orders/${orderId}`, {
    headers: { Authorization: `Bearer ${token}` },
    data: { customerName: 'Jane Smith' }
});

// DELETE
const r = await request.delete(`/orders/${orderId}`, {
    headers: { Authorization: `Bearer ${token}` }
});""")

    h2(doc, '9.6 Reading the Response')
    code(doc, """\
r.status()                   // number:  200
r.ok()                       // boolean: true if status < 400
await r.json()               // parsed object or array (must use await!)
await r.text()               // raw string body
r.headers()['content-type']  // header value""")

    h2(doc, '9.7 Assertions with expect()')
    code(doc, """\
expect(response.status()).toBe(201);                   // exact equality
expect(response.status()).not.toBe(400);               // negation
expect(books.length).toBeGreaterThan(0);               // greater than zero
expect(books.length).toBeLessThanOrEqualTo(2);         // at most 2
expect(body).toHaveProperty('orderId');                // key exists in object
expect(body.created).toBe(true);                       // boolean value
expect(body.customerName).toBe('Jane Smith');          // string equality
expect(Array.isArray(books)).toBe(true);               // type check""")

    h2(doc, '9.8 playwright.config.ts — The Control Centre')
    code(doc, """\
export default defineConfig({
  testDir: './tests',
  fullyParallel: true,     // run all tests in parallel by default
  reporter: [
    ['html', { open: 'never' }],   // generates playwright-report/index.html
    ['list'],                       // prints live results in terminal
  ],
  use: {
    baseURL: 'https://simple-books-api.click',
    // ↑ This means request.get('/books') automatically becomes
    //   request.get('https://simple-books-api.click/books')
    extraHTTPHeaders: {
      'Content-Type': 'application/json',
      'Accept': 'application/json',
    },
  },
});""")

    h2(doc, '9.9 Serial vs Parallel Tests')
    table(doc,
        ['', 'Regular test.describe', 'test.describe.serial'],
        [
            ['Execution order', 'Tests run in parallel (simultaneously)', 'Tests run in strict sequential order'],
            ['Speed',           'Faster',                                  'Slower (one at a time)'],
            ['Use when',        'Tests are independent of each other',     'Test 2 depends on result of Test 1'],
            ['Orders tests?',   'WRONG — test 2 starts before test 1 ends','CORRECT'],
        ],
        col_widths=[1.4, 2.5, 2.5]
    )
    code(doc, """\
// Use this for orders tests:
test.describe.serial('Orders API', () => {
  let token: string;
  let orderId: string;

  test.beforeAll(async ({ request }) => {
    token = await getAccessToken(request);  // runs once before all tests below
  });

  test('POST /orders', async ({ request }) => {
    const r = await request.post('/orders', { ... });
    orderId = (await r.json()).orderId;  // saved for next test
  });

  test('PATCH /orders/:id', async ({ request }) => {
    // orderId is available here because the previous test ran first
    await request.patch(`/orders/${orderId}`, { ... });
  });
});""")

    h2(doc, '9.10 How to Run — TypeScript Playwright')
    table(doc,
        ['Command', 'What It Does'],
        [
            ['npm install',                          'Install dependencies'],
            ['npx playwright test',                  'Run all tests'],
            ['npx playwright test tests/api',        'Run API tests folder only'],
            ['npx playwright test tests/api/orders', 'Run orders spec only'],
            ['npx playwright show-report',           'Open HTML visual report'],
            ['npx playwright test --debug',          'Interactive debugger mode'],
        ],
        col_widths=[3.2, 3.2]
    )

    page_break(doc)


# ── CHAPTER 10 ────────────────────────────────────────────────────────────────

def chapter_10(doc):
    h1(doc, 'Chapter 10 — Test Design Principles')

    h2(doc, '10.1 The Testing Pyramid')
    code(doc, """\
           /\\
          /  \\     End-to-End / UI Tests
         / E2E \\   Slow, expensive, brittle — write few
        /________\\
       /          \\
      / Integration \\  API Tests  ← We are here
     /   (API Tests)  \\ Fast, stable, great ROI
    /__________________\\
   /                    \\
  /    Unit  Tests        \\ Fastest, cheapest — write the most
 /________________________\\""")
    body(doc, 'API tests are the sweet spot: they test full request/response cycles '
              '(including business logic and database) without the overhead and '
              'fragility of browser automation.')

    h2(doc, '10.2 FIRST Principles')
    table(doc,
        ['Letter', 'Principle', 'Meaning'],
        [
            ['F', 'Fast',          'Tests should run in seconds, not minutes'],
            ['I', 'Independent',   'Tests should not depend on each other*'],
            ['R', 'Repeatable',    'Same result every time, on any machine'],
            ['S', 'Self-Validating','Pass or fail automatically — no manual review'],
            ['T', 'Timely',        'Written at the same time as the code being tested'],
        ],
        col_widths=[0.4, 1.5, 4.5]
    )
    note_box(doc, '* Exception: CRUD lifecycle tests (Create → Read → Update → Delete) '
                  'have intentional ordering. This is acceptable when documenting the dependency clearly.')

    h2(doc, '10.3 Good Test Naming')
    body(doc, 'A good test name tells you EXACTLY what is being tested and what is expected:')
    code(doc, """\
# BAD names — too vague:
def test1():
def test_order():
def test_api():

# GOOD names — self-documenting:
def test_create_order_returns_201_with_order_id():
def test_create_order_without_token_returns_401():
def test_get_non_existent_book_returns_404():
def test_filter_books_by_fiction_returns_only_fiction_type():""")

    h2(doc, '10.4 Arrange — Act — Assert (AAA Pattern)')
    body(doc, 'Every test should have three clearly separated sections:')
    code(doc, """\
def test_create_order(self, base_url, auth_headers):

    # ── ARRANGE: set up the data you need ──────────────────
    payload = {
        "bookId": 1,
        "customerName": "John Doe"
    }

    # ── ACT: perform the action being tested ───────────────
    response = requests.post(
        f"{base_url}/orders",
        json=payload,
        headers=auth_headers
    )

    # ── ASSERT: check the result ───────────────────────────
    assert response.status_code == 201
    assert response.json()["created"] is True
    assert "orderId" in response.json()""")

    h2(doc, '10.5 Test Coverage Checklist')
    body(doc, 'Ask these questions for every endpoint:')
    numbered(doc, 'Does the happy path (valid data, valid auth) return the correct status and body?')
    numbered(doc, 'What happens with missing required fields? → expect 400')
    numbered(doc, 'What happens with an invalid resource ID? → expect 404')
    numbered(doc, 'What happens without an auth token? → expect 401')
    numbered(doc, 'What happens with invalid query parameter values? → expect 400')
    numbered(doc, 'After a PATCH, does a GET confirm the update was actually saved?')
    numbered(doc, 'After a DELETE, does a GET confirm the resource is gone (404)?')

    page_break(doc)


# ── CHAPTER 11 — CHEAT SHEETS ─────────────────────────────────────────────────

def chapter_11(doc):
    h1(doc, 'Chapter 11 — Cheat Sheets & Quick Reference')

    h2(doc, '11.1 HTTP Status Code Quick Reference')
    code(doc, """\
200 OK             → GET succeeded, body has the data
201 Created        → POST succeeded, new resource created
204 No Content     → PATCH/DELETE succeeded, no body
400 Bad Request    → You sent bad data (missing field, wrong type)
401 Unauthorized   → No token or invalid token
403 Forbidden      → Token valid but no permission for this resource
404 Not Found      → Resource doesn't exist at that URL
409 Conflict       → Duplicate (same email registered twice)
500 Server Error   → Bug on the server side — not your fault""")

    h2(doc, '11.2 The Same Test in All 3 Languages')
    body(doc, 'Scenario: Create an order and verify the 201 response.')

    h3(doc, 'Java — REST Assured')
    code(doc, """\
@Test
public void testCreateOrder() {
    Map<String, Object> body = new HashMap<>();
    body.put("bookId", 1);
    body.put("customerName", "John Doe");

    Response response = RestAssured.given()
            .spec(requestSpec)
            .header("Authorization", "Bearer " + token)
            .body(body)
            .when().post("/orders")
            .then().statusCode(201)
                   .body("created", equalTo(true))
                   .extract().response();

    orderId = response.jsonPath().getString("orderId");
    assertNotNull(orderId);
}""")

    h3(doc, 'Python — pytest + requests')
    code(doc, """\
def test_create_order(self, base_url, auth_headers):
    payload = {"bookId": 1, "customerName": "John Doe"}

    response = requests.post(
        f"{base_url}/orders",
        json=payload,
        headers=auth_headers
    )

    assert response.status_code == 201
    data = response.json()
    assert data["created"] is True
    assert "orderId" in data

    TestOrders.order_id = data["orderId"]""")

    h3(doc, 'TypeScript — Playwright')
    code(doc, """\
test('POST /orders — create order', async ({ request }) => {
    const response = await request.post('/orders', {
        headers: { Authorization: `Bearer ${token}` },
        data: { bookId: 1, customerName: 'John Doe' },
    });

    expect(response.status()).toBe(201);
    const body = await response.json();
    expect(body.created).toBe(true);
    expect(body).toHaveProperty('orderId');

    orderId = body.orderId;
});""")

    h2(doc, '11.3 REST Assured (Java) — Common Patterns')
    code(doc, """\
// GET with query param
given().spec(spec).queryParam("type", "fiction")
       .when().get("/books").then().statusCode(200);

// GET with path param
given().spec(spec).pathParam("id", orderId)
       .when().get("/orders/{id}").then().statusCode(200);

// POST with auth + body
given().spec(spec)
       .header("Authorization", "Bearer " + token).body(map)
       .when().post("/orders").then().statusCode(201);

// PATCH (204 = no body, don't read response body)
given().spec(spec)
       .header("Authorization", "Bearer " + token).body(updateMap)
       .when().patch("/orders/{id}").then().statusCode(204);

// DELETE
given().spec(spec)
       .header("Authorization", "Bearer " + token)
       .when().delete("/orders/{id}").then().statusCode(204);""")

    h2(doc, '11.4 requests (Python) — Common Patterns')
    code(doc, """\
# GET with params
response = requests.get(url, params={"type": "fiction"})

# POST with JSON body
response = requests.post(url, json={"bookId": 1}, headers=headers)

# PATCH
response = requests.patch(url, json={"customerName": "Jane"}, headers=headers)

# DELETE
response = requests.delete(url, headers=headers)

# Reading responses
response.status_code          # int
response.json()               # dict or list
response.json()["key"]        # access a field directly
response.json().get("key")    # safer: returns None if key missing
response.text                 # raw string body""")

    h2(doc, '11.5 Playwright (TypeScript) — Common Patterns')
    code(doc, """\
// GET with params → /books?type=fiction
await request.get('/books', { params: { type: 'fiction' } });

// POST with body and auth
await request.post('/orders', {
    headers: { Authorization: `Bearer ${token}` },
    data: { bookId: 1, customerName: 'John' }
});

// PATCH / DELETE
await request.patch(`/orders/${id}`, {
    headers: { Authorization: `Bearer ${token}` },
    data: { customerName: 'Jane' }
});
await request.delete(`/orders/${id}`, {
    headers: { Authorization: `Bearer ${token}` }
});

// Assertions
expect(r.status()).toBe(200);
expect(body).toHaveProperty('orderId');
expect(array.length).toBeGreaterThan(0);""")

    h2(doc, '11.6 Full Command Reference')
    table(doc,
        ['Tool', 'Command', 'Purpose'],
        [
            ['Java',       'mvn test',                               'Run all tests'],
            ['Java',       'mvn test -Dtest=BooksTest',             'Run one class'],
            ['Python',     'pip install -r requirements.txt',       'Install dependencies'],
            ['Python',     'pytest -v',                             'Run all, verbose'],
            ['Python',     'pytest --html=report.html',            'Generate HTML report'],
            ['Python',     'pytest -x',                             'Stop on first failure'],
            ['Playwright', 'npm install',                           'Install dependencies'],
            ['Playwright', 'npx playwright test tests/api',        'Run API tests'],
            ['Playwright', 'npx playwright show-report',           'Open visual report'],
            ['Playwright', 'npx playwright test --debug',         'Interactive debugger'],
        ],
        col_widths=[1.1, 2.8, 2.5]
    )

    h2(doc, '11.7 Framework Comparison Summary')
    table(doc,
        ['Feature', 'Java REST Assured', 'Python pytest', 'TypeScript Playwright'],
        [
            ['HTTP Client',     'REST Assured',             'requests library',        'Built-in request fixture'],
            ['Test Framework',  'TestNG',                   'pytest',                  'Playwright Test'],
            ['Setup/Teardown',  '@BeforeSuite, @BeforeClass','@pytest.fixture',         'test.beforeAll'],
            ['Assertions',      'Hamcrest matchers',        'assert statement',        'expect() matchers'],
            ['Token Caching',   'Static class field',       'scope="session" fixture', 'Module-level variable'],
            ['Parallel Tests',  'testng.xml parallel=',     'pytest-xdist plugin',     'fullyParallel: true'],
            ['Run Command',     'mvn test',                 'pytest',                  'npx playwright test'],
            ['Report',          'TestNG HTML report',       'pytest-html plugin',      'Built-in HTML report'],
        ],
        col_widths=[1.5, 1.7, 1.7, 1.7]
    )


# ── MAIN ──────────────────────────────────────────────────────────────────────

def main():
    doc = Document()

    # Set default page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Cover Page
    cover_page(doc)

    # Table of Contents page
    h1(doc, 'Table of Contents')
    body(doc, 'Right-click the table of contents below and select "Update Field" '
              'to refresh page numbers after opening the document.')
    add_toc(doc)
    page_break(doc)

    # All chapters
    chapter_1(doc)
    chapter_2(doc)
    chapter_3(doc)
    chapter_4(doc)
    chapter_5(doc)
    chapter_6(doc)
    chapter_7(doc)
    chapter_8(doc)
    chapter_9(doc)
    chapter_10(doc)
    chapter_11(doc)

    # Save
    output_path = r'd:\Downloads\simple-books-api-tests\API_Automation_Learning_Guide.docx'
    doc.save(output_path)
    print(f"\nDONE. Document saved to:\n  {output_path}\n")
    print("Tips:")
    print("  - Use the Navigation Pane (View > Navigation Pane) to jump between chapters.")
    print("  - Right-click the Table of Contents > Update Field to refresh page numbers.")
    print("  - Code blocks have a light-gray background.")
    print("  - Tables use alternating row colors for easy reading.\n")


if __name__ == '__main__':
    main()
